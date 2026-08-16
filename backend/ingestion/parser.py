"""Document parsing: extract text content from various file formats."""

from pathlib import Path

from backend.core.exceptions import DocumentParsingError, UnsupportedFileTypeError
from backend.core.logging import get_logger
from backend.models.documents import SUPPORTED_EXTENSIONS, FileType

logger = get_logger(__name__)


class ParsedDocument:
    """Result of parsing a document file.

    Attributes:
        text: The full extracted text content.
        pages: Per-page text content (list index = page index). Empty for non-paged formats.
        filename: Original filename.
        file_type: Detected file type.
    """

    def __init__(self, text: str, pages: list[str], filename: str, file_type: FileType) -> None:
        self.text = text
        self.pages = pages
        self.filename = filename
        self.file_type = file_type


class DocumentParser:
    """Parses uploaded documents into plain text.

    Supports PDF, DOCX, TXT, Markdown, and CSV formats.
    Each format is handled by a dedicated private method.
    """

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a document file and extract its text content.

        Args:
            file_path: Path to the file on disk.

        Returns:
            ParsedDocument containing the extracted text and page information.

        Raises:
            UnsupportedFileTypeError: If the file extension is not supported.
            DocumentParsingError: If parsing fails for any reason.
        """
        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(file_path.name, suffix)

        file_type = SUPPORTED_EXTENSIONS[suffix]
        logger.info("Parsing document", filename=file_path.name, file_type=file_type)

        try:
            if file_type == FileType.PDF:
                return self._parse_pdf(file_path, file_type)
            elif file_type == FileType.DOCX:
                return self._parse_docx(file_path, file_type)
            elif file_type in (FileType.TXT, FileType.MARKDOWN, FileType.CSV):
                return self._parse_text(file_path, file_type)
            else:
                raise UnsupportedFileTypeError(file_path.name, suffix)
        except (UnsupportedFileTypeError, DocumentParsingError):
            raise
        except Exception as exc:
            raise DocumentParsingError(file_path.name, str(exc)) from exc

    def _parse_pdf(self, file_path: Path, file_type: FileType) -> ParsedDocument:
        """Extract text from a PDF file, page by page.

        Args:
            file_path: Path to the PDF file.
            file_type: The file type enum value.

        Returns:
            ParsedDocument with per-page text content.
        """
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        pages: list[str] = []

        for page in reader.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text.strip())

        full_text = "\n\n".join(pages)
        logger.info("Parsed PDF", filename=file_path.name, page_count=len(pages))
        return ParsedDocument(text=full_text, pages=pages, filename=file_path.name, file_type=file_type)

    def _parse_docx(self, file_path: Path, file_type: FileType) -> ParsedDocument:
        """Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file.
            file_type: The file type enum value.

        Returns:
            ParsedDocument with full text (no page-level separation for DOCX).
        """
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        paragraphs: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)
        logger.info("Parsed DOCX", filename=file_path.name, paragraph_count=len(paragraphs))
        return ParsedDocument(text=full_text, pages=[], filename=file_path.name, file_type=file_type)

    def _parse_text(self, file_path: Path, file_type: FileType) -> ParsedDocument:
        """Read plain text, Markdown, or CSV files.

        Args:
            file_path: Path to the text file.
            file_type: The file type enum value.

        Returns:
            ParsedDocument with the raw text content.
        """
        content = file_path.read_text(encoding="utf-8")
        logger.info("Parsed text file", filename=file_path.name, char_count=len(content))
        return ParsedDocument(text=content, pages=[], filename=file_path.name, file_type=file_type)
